from new_paper import *
import util
import pandas as pd
from docx import Document
from docx.shared import Inches,Pt,RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


def make_header(document, paperConfig, is_mcq=False, is_half=False):
    q_code = document.add_heading('Question Paper Code:' + str(paperConfig['Question Paper Code']),level = 6)
    q_code.alignment = WD_ALIGN_PARAGRAPH.RIGHT


    heading = document.add_heading(level = 3)
    title = heading.add_run('SRI DEVRAJ URS ACADEMY OF HIGHER EDUCATION & RESEARCH')
    title.underline = True
    title.font.color.rgb = RGBColor(0,0,0)
    temp = heading.add_run('''\n(A DEEMED TO BE UNIVERSITY)''')
    temp.italic = True
    temp.font.color.rgb = RGBColor(0x00,0x00,0x00)
    exam = heading.add_run('''\n'''+ str(paperConfig['exam name']))
    exam.font.color.rgb = RGBColor(0,0,0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    rules = document.add_paragraph()
    
    if is_mcq == True:
        time = rules.add_run('Time:' + "30 minutes" + '\t\t\t\t\t\t\tMax Marks:' + "20")
    else:
        if is_half:
            time = rules.add_run('Time:' + str(paperConfig['Max duration']) + '\t\t\t\t\t\t\tMax Marks:100') 
        else:
            time = rules.add_run('Time:' + str(paperConfig['Max duration']) + '\t\t\t\t\t\t\tMax Marks:' + paperConfig['Max marks'])
    #time = rules.add_run('Time: 3 hours')
    time.alignment = WD_ALIGN_PARAGRAPH.LEFT
    time.bold = True
    time.font.size = Pt(11.5)
    
    subject = document.add_heading(level = 3)
    sub = subject.add_run(paperConfig['subject'])
    sub.font.color.rgb = RGBColor(0,0,0)
    subject.alignment = WD_ALIGN_PARAGRAPH.CENTER
   
    important_notes = document.add_paragraph()
    
    if is_mcq == False:
        temp = paperConfig['ins'].split('.')
        res = '\n'
        for sent in temp:
            res = res + sent + '\n'

        i = important_notes.add_run(res)

        i.italic = True
        i.font.size = Pt(8)
        important_notes.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    return document


def make_paper(document, paper, is_mcq=False, with_options=False, is_half=False):
    cat1  = document.add_heading(level = 4)
    tabs = ""
    if with_options:
        topic = cat1.add_run('LONG ESSAY ( Answer any 2 )')
        tabs = "\t\t\t\t\t\t" 
    else:
        topic = cat1.add_run('LONG ESSAY')
        tabs = "\t\t\t\t\t\t\t\t"
    topic.font.color.rgb = RGBColor(0,0,0)
    topic.font.size = Pt(10.5)
    topic.underline = True
    marks = 20
    if is_half: marks = 10
    topic = cat1.add_run(f'{tabs}{int(marks/10)} X 10 = {marks} Marks')
    topic.font.color.rgb = RGBColor(0,0,0)
    topic.font.size = Pt(10.5)


    result = paper[paper["type"] == 'le']
    for index, row in result.iterrows():
        p = document.add_paragraph(style = 'ListNumber') 
        q = p.add_run(row['question'])
        q.font.size = Pt(11)

    cat2  = document.add_heading(level = 4)
    tabs = ""
    if with_options:
        topic = cat2.add_run('SHORT ESSAY ( Answer any 10 )')
        tabs = "\t\t\t\t\t\t"
    else:
        topic = cat2.add_run('SHORT ESSAY')
        tabs = "\t\t\t\t\t\t\t\t"

    topic.font.color.rgb = RGBColor(0,0,0)
    topic.font.size = Pt(10.5)
    topic.underline = True
    if not is_mcq:
        marks = 50
        if is_half: marks = 25
        topic = cat2.add_run(f'{tabs}{int(marks/5)} X 5 = {marks} Marks')
    else:
        marks = 30
        if is_half: marks = 15
        topic = cat2.add_run(f'{tabs}{int(marks/5)} X 5 = {marks} Marks')
    topic.font.color.rgb = RGBColor(0,0,0)
    topic.font.size = Pt(10.5)


    result = paper[paper["type"] == 'se']
    for index, row in result.iterrows():
        p = document.add_paragraph(style = 'ListNumber') 
        q = p.add_run(row['question'])
        q.font.size = Pt(11)


    cat3  = document.add_heading(level = 4)
    tabs = ""
    if with_options:
        topic =  cat3.add_run('SHORT ANSWERS ( No choices )')
        tabs = "\t\t\t\t\t\t"
    else:
        topic = cat3.add_run('SHORT ANSWERS')
        tabs = "\t\t\t\t\t\t\t"
    topic.font.color.rgb = RGBColor(0,0,0)
    topic.font.size = Pt(10.5)
    topic.underline = True
    marks = 30
    if is_half: marks = 15
    topic = cat3.add_run(f'{tabs}{int(marks/3)} X 3 = {marks} Marks')
    topic.font.color.rgb = RGBColor(0,0,0)
    topic.font.size = Pt(10.5)

    result = paper[paper["type"] == 'sa']
    for index, row in result.iterrows():
        p = document.add_paragraph(style = 'ListNumber') 
        q = p.add_run(row['question'])
        q.font.size = Pt(11)

    return document

def add_section(document, section_name):
    subject = document.add_heading(level = 3)
    sub = subject.add_run(section_name)
    sub.font.color.rgb = RGBColor(0,0,0)
    subject.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return document


def make_mcq_paper(document, paper):
    cat1  = document.add_heading(level = 4)
    topic = cat1.add_run('Multiple Choice Questions')
    topic.font.color.rgb = RGBColor(0,0,0)
    topic.font.size = Pt(10.5)
    topic.underline = True
    topic = cat1.add_run('\t\t\t\t\t\t1 X 20 = 20 Marks')
    topic.font.color.rgb = RGBColor(0,0,0)
    topic.font.size = Pt(10.5)

    result = paper[paper["type"] == 'mcq']
    for index, row in result.iterrows():
        p = document.add_paragraph(style = 'ListNumber') 
        q = p.add_run(row['question'])
        q.add_break()
        q.add_tab()
        q = p.add_run("a) ")
        q = p.add_run(row['option1'])
        q.add_break()
        q.add_tab()
        q = p.add_run("b) ")
        q = p.add_run(row['option2'])
        q.add_break()
        q.add_tab()
        q = p.add_run("c) ")
        q = p.add_run(row['option3'])
        q.add_break()
        q.add_tab()
        q = p.add_run("d) ")
        q = p.add_run(row['option4'])
        q.add_break()
        q.font.size = Pt(11)

    return document



if __name__=="__main__":
    config = util.read_csv_python("format.csv", ["key", "value"])
    config_dict = pd.Series(config.value.values,index=config.key.values).to_dict()
    print(config_dict)

    document = Document()
    document1 = make_header(document, config_dict)

    document.save("testheader.docx")

    df = pd.read_csv("paper1.csv")
    document = make_paper(document1, df)

    document.save("test.docx")

    document = Document()
    document1 = make_header(document, config_dict, True)

    document = make_mcq_paper(document1, df)

    document.save("test2.docx")



