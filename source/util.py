import random
import pandas as pd
from collections import OrderedDict
import numpy as np
from docx import Document
from docx.shared import Inches,Pt,RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def flip(p):
    return (random.random() < p)

#make paper blueprint
def fillBuckets(topicSplit, qContrib, qCand, topicBreakup, topicBreakup_cache):
    qBuckets = {}
    inverseBuckets = {}
    accumulatedPaperTotal = 0
    for category in qContrib:
        qBuckets[category] = []
    items = list(reversed(topicSplit.items()))
    #long essay questions
    a,b = items[0]
    c,d = items[1]

    qBuckets['le'].append(a)
    qBuckets['le'].append(c)
    
    topicSplit[a] = b - 0.1
    topicSplit[c] = d - 0.1
    topicSplit = OrderedDict(sorted(topicSplit.items(), key = lambda t: t[1]))
    
    accumulatedPaperTotal = 0.2
    #qCand['le'] = 0
    
    for i in range(10):
        topicSplit = OrderedDict(sorted(topicSplit.items(), key = lambda t: t[1]))
        items = list(reversed(topicSplit.items()))
        a,b = items[0]
        if(b > 0.08):
            qBuckets['se'].append(a)
            qBuckets['sa'].append(a)
            topicSplit[a] = b - 0.08
        elif(b > 0.05):
            qBuckets['se'].append(a)
            topicSplit[a] = b - 0.05
        else:
            qBuckets['sa'].append(a)
            topicSplit[a] = b - 0.03
    '''
    for bck in qBuckets:
        print(len(qBuckets[bck]))
    '''
    
    while(len(qBuckets['se']) < 10):
        topicSplit = OrderedDict(sorted(topicSplit.items(), key = lambda t: t[1]))
        items = list(reversed(topicSplit.items()))
        a,b = items[0]
        qBuckets['se'].append(a)
        topicSplit[a] = b - 0.05
    
    while(len(qBuckets['sa']) < 10):
        topicSplit = OrderedDict(sorted(topicSplit.items(), key = lambda t: t[1]))
        items = list(reversed(topicSplit.items()))
        a,b = items[0]
        qBuckets['sa'].append(a)
        topicSplit[a] = b - 0.05

    data = {'category' : [],'difficulty' : [],'topic' : []} 
    for category in qBuckets:
        for topic in qBuckets[category]:
            data['category'].append(category)
            data['topic'].append(topic)
            data['difficulty'].append('low')

    df = pd.DataFrame(data)
    almost_df = populateDifficulty(df,qContrib, topicBreakup, topicBreakup_cache)
    working_df = almost_df.copy(deep=True)
    for a,b in almost_df.iterrows():
        candidates = b['topic'].split('/')
        chosen_one = random.choice(candidates)
        working_df.set_value(a,'topic',chosen_one)
        
    return working_df

def fillBucketsWithOptions(topicSplit, qContrib, qCand, topicBreakup, topicBreakup_cache):
    qBuckets = {}
    inverseBuckets = {}
    accumulatedPaperTotal = 0
    for category in qContrib:
        qBuckets[category] = []
    items = list(reversed(topicSplit.items()))
    #long essay questions
    a,b = items[0]
    c,d = items[1]
    e,f = items[2]

    qBuckets['le'].append(a)
    qBuckets['le'].append(c)
    qBuckets['le'].append(e)
    
    topicSplit[a] = b - 0.1
    topicSplit[c] = d - 0.1
    topicSplit[e] = f - 0.1
    topicSplit = OrderedDict(sorted(topicSplit.items(), key = lambda t: t[1]))
    
    accumulatedPaperTotal = 0.2
    #qCand['le'] = 0
    
    for i in range(10):
        topicSplit = OrderedDict(sorted(topicSplit.items(), key = lambda t: t[1]))
        items = list(reversed(topicSplit.items()))
        a,b = items[0]
        if(b > 0.08):
            qBuckets['se'].append(a)
            qBuckets['sa'].append(a)
            topicSplit[a] = b - 0.08
        elif(b > 0.05):
            qBuckets['se'].append(a)
            topicSplit[a] = b - 0.05
        else:
            qBuckets['sa'].append(a)
            topicSplit[a] = b - 0.03
    '''
    for bck in qBuckets:
        print(len(qBuckets[bck]))
    '''

    while(len(qBuckets['se']) < 10):
        topicSplit = OrderedDict(sorted(topicSplit.items(), key = lambda t: t[1]))
        items = list(reversed(topicSplit.items()))
        a,b = items[0]
        qBuckets['se'].append(a)
        topicSplit[a] = b - 0.05
    
    while(len(qBuckets['sa']) < 10):
        topicSplit = OrderedDict(sorted(topicSplit.items(), key = lambda t: t[1]))
        items = list(reversed(topicSplit.items()))
        a,b = items[0]
        qBuckets['sa'].append(a)
        topicSplit[a] = b - 0.05
    
    topicSplit = OrderedDict(sorted(topicSplit.items(), key = lambda t: t[1]))
    items = list(reversed(topicSplit.items()))
    a,b = items[0]
    qBuckets['se'].append(a)
    topicSplit[a] = b - 0.05

    topicSplit = OrderedDict(sorted(topicSplit.items(), key = lambda t: t[1]))
    items = list(reversed(topicSplit.items()))
    a,b = items[0]
    qBuckets['se'].append(a)
    topicSplit[a] = b - 0.05

    dataWithOptions = {'category' : [],'difficulty' : [],'topic' : []}
    for category in qBuckets:
        for topic in qBuckets[category]:
            dataWithOptions['category'].append(category)
            dataWithOptions['topic'].append(topic)
            dataWithOptions['difficulty'].append('low')
    dfWithOptions = pd.DataFrame(dataWithOptions)


    almost_dfWithOptions = populateDifficulty(dfWithOptions, qContrib, topicBreakup, topicBreakup_cache)
    working_dfWithOptions = almost_dfWithOptions.copy(deep=True)
    for a,b in almost_dfWithOptions.iterrows():
        candidates = b['topic'].split('/')
        chosen_one = random.choice(candidates)
        working_dfWithOptions.set_value(a,'topic',chosen_one)
    return working_dfWithOptions

def populateDifficulty(df, qContrib, topicBreakup, topicBreakup_cache):
    working_df = df.copy(deep=True)
    difficultySplit = {}
    topicWeights = {}
    difficultyBreakup = {'low' : .5, 'medium' : .3, 'high' : .2}
    accumulatedDifficulty =  { 'low' : 0 , 'medium' : 0, 'high' : 0}
    for topic in topicBreakup:
        parcel = { 'low' : 0 , 'medium' : 0, 'high' : 0}
        difficultySplit[topic] = parcel
        topicWeights[topic] = {}
        for category in qContrib:
            topicWeights[topic][category] = qContrib[category] / topicBreakup_cache[topic]
    
    le_count = 0
    for index,row in df.iterrows():
        if le_count == 0 and row['category'] == 'le':
            working_df.set_value(index, 'difficulty', 'low')
            le_count += 1
            difficultySplit[row['topic']]['low'] += topicWeights[row['topic']]['le']
            accumulatedDifficulty['low'] += qContrib['le']

        elif row['category'] == 'le' and flip(0.0):
            working_df.set_value(index, 'difficulty', 'low')
            difficultySplit[row['topic']]['low'] += topicWeights[row['topic']]['le']
            accumulatedDifficulty['low'] += qContrib['le']

        elif row['category'] == 'le':
            working_df.set_value(index, 'difficulty', 'medium')
            difficultySplit[row['topic']]['medium'] += topicWeights[row['topic']]['le']
            accumulatedDifficulty['medium'] += qContrib['le']


        elif row['category'] == 'se' and difficultySplit[row['topic']]['high'] < difficultyBreakup['high'] and topicWeights[row['topic']]['se'] < difficultyBreakup['high'] and accumulatedDifficulty['high'] < difficultyBreakup['high'] and flip(0.55 - difficultySplit[row['topic']]['high'] ):
            working_df.set_value(index, 'difficulty', 'high')
            difficultySplit[row['topic']]['high'] += topicWeights[row['topic']]['se']
            accumulatedDifficulty['high'] += qContrib['se']
        
        elif row['category'] == 'se' and accumulatedDifficulty['high'] + topicWeights[row['topic']]['sa'] / 2 < difficultyBreakup['high'] and flip(0.3):
            working_df.set_value(index, 'difficulty', 'high')
            difficultySplit[row['topic']]['high'] += topicWeights[row['topic']]['se']
            accumulatedDifficulty['high'] += qContrib['se']

        elif row['category'] == 'se' and accumulatedDifficulty['high'] + 0.05 < difficultyBreakup['high'] and flip(0.3):
            working_df.set_value(index, 'difficulty', 'high')
            difficultySplit[row['topic']]['high'] += topicWeights[row['topic']]['se']
            accumulatedDifficulty['high'] += qContrib['se']
        
        elif row['category'] == 'se' and difficultySplit[row['topic']]['medium'] < difficultyBreakup['medium'] and topicWeights[row['topic']]['se'] < difficultyBreakup['medium'] and accumulatedDifficulty['medium'] + 0.05 < difficultyBreakup['medium'] and flip(0.8 - difficultySplit[row['topic']]['medium']):
            working_df.set_value(index, 'difficulty', 'medium')
            difficultySplit[row['topic']]['medium'] += topicWeights[row['topic']]['se']
            accumulatedDifficulty['medium'] += qContrib['se']

        elif row['category'] == 'se' and difficultySplit[row['topic']]['medium'] < difficultyBreakup['medium'] and accumulatedDifficulty['medium'] + 0.05 < difficultyBreakup['medium'] and flip(0.8 - difficultySplit[row['topic']]['medium']):
            working_df.set_value(index, 'difficulty', 'medium')
            difficultySplit[row['topic']]['medium'] += topicWeights[row['topic']]['se']
            accumulatedDifficulty['medium'] += qContrib['se']

        elif row['category'] == 'se' :
            working_df.set_value(index, 'difficulty', 'low')
            difficultySplit[row['topic']]['low'] += topicWeights[row['topic']]['se']
            accumulatedDifficulty['low'] += qContrib['se']

        elif row['category'] == 'sa' and difficultySplit[row['topic']]['high'] < difficultyBreakup['high'] and topicWeights[row['topic']]['sa'] + difficultySplit[row['topic']]['high'] < difficultyBreakup['high'] and accumulatedDifficulty['high'] < difficultyBreakup['high'] and flip(0.45 - difficultySplit[row['topic']]['high'] ):
            working_df.set_value(index, 'difficulty', 'high')
            difficultySplit[row['topic']]['high'] += topicWeights[row['topic']]['sa']
            accumulatedDifficulty['high'] += qContrib['sa']
        
        elif row['category'] == 'sa' and accumulatedDifficulty['high']< difficultyBreakup['high'] and flip(0.5):
            working_df.set_value(index, 'difficulty', 'high')
            difficultySplit[row['topic']]['high'] += topicWeights[row['topic']]['sa']
            accumulatedDifficulty['high'] += qContrib['sa']

        elif row['category'] == 'sa' and difficultySplit[row['topic']]['medium'] < difficultyBreakup['medium'] and topicWeights[row['topic']]['sa'] + difficultySplit[row['topic']]['medium'] < difficultyBreakup['medium'] and accumulatedDifficulty['medium'] < difficultyBreakup['medium'] and flip(1 - difficultySplit[row['topic']]['medium'] ):
            working_df.set_value(index, 'difficulty', 'medium')
            difficultySplit[row['topic']]['medium'] += topicWeights[row['topic']]['sa']
            accumulatedDifficulty['medium'] += qContrib['sa']

        elif row['category'] == 'sa' and difficultySplit[row['topic']]['medium'] == 0 :
            working_df.set_value(index, 'difficulty', 'medium')
            difficultySplit[row['topic']]['medium'] += topicWeights[row['topic']]['sa']
            accumulatedDifficulty['medium'] += qContrib['sa']

        else:
            working_df.set_value(index, 'difficulty', 'low')
            difficultySplit[row['topic']]['low'] += topicWeights[row['topic']]['sa']
            accumulatedDifficulty['low'] += qContrib['sa']
        
        '''candidates = row['topic'].split(' / ')
        print(candidates)

    print(working_df)
    working_df.to_csv("generatedBluePrint.csv")
    f = pd.read_csv('generatedBluePrint.csv', index_col = 0)
        '''

    '''print("\ncheck \n")
    print(f,"\n")

    print(difficultySplit)
    '''
    print('accumulated_difficulty : \n', accumulatedDifficulty)
    return(working_df)

def setSubTopicMask(df, sub_topic, question):
    #print(sub_topic)
    df['temp'] = np.where(df['sub_topic'] == sub_topic,1,df['mask'])
    df['temp1'] = np.where(df['question'] == question,1,df['picked'])
    df = df.drop('mask',axis = 1)
    df = df.drop('picked',axis = 1)
    df.rename(columns={'temp':'mask'},inplace = True)
    df.rename(columns={'temp1':'picked'},inplace = True)
    #print(df)
    return df

def pickQuestion(df, topic, category, difficulty):
    result = df[(df['topic'] == topic) & (df['difficulty']==difficulty) & (df['category']==category) & (df['mask'] != 1)]
    if not result.empty:
        result = result.sample(n=1)
        #print(" HIT 1\n")
    else:
        result = df[(df['topic'] == topic) & (df['category']==category) & (df['mask'] != 1)]
        if not result.empty:
            result = result.sample(n=1)
            #print(" HIT 2\n")
        else:
            result = df[(df['topic'] == topic) & (df['category']==category) & (df['picked'] != 1)]
            if not result.empty:
                result = result.sample(n=1)
                #print(" HIT 3\n")
            else:
                result = df[(df['difficulty']==difficulty) & (df['category']==category)   & (df['picked'] != 1)]
                if not result.empty:
                    result = result.sample(n=1)
                    #print(" HIT 4 " + topic + "\n")
                else:
                    result = df[(df['category']==category)  & (df['picked'] != 1)]
                    if not result.empty:
                        result = result.sample(n=1)
                        #print("WTFFFFF\n\n\n\n\n\n")

            

    return result

def populateBluePrint(bp, qb):
    count = 0
    for a,b in bp.iterrows():
        result = pickQuestion(qb, b['topic'], b['category'], b['difficulty'])
        if not result.empty:
            count+=1
            result = result.reset_index(drop=True)
            b['question'] = result['question'][0]
            b['sub_topic'] = result['sub_topic'][0]
            b['difficulty'] = result['difficulty'][0]
            qb = setSubTopicMask(qb, b['sub_topic'], b['question'])
        else:
            print("no question satisfying constraints\n\n\n")
    bp.to_csv('QPaper.csv')
    return bp

def setPaper_no_options(df, paperConfig, name):
    document = Document()

    q_code = document.add_heading('Question Paper Code:' + str(paperConfig['Question Paper Code']),level = 6)
    q_code.alignment = WD_ALIGN_PARAGRAPH.RIGHT


    heading = document.add_heading(level = 3)
    title = heading.add_run('SRI DEVRAJ URS ACADEMY OF HIGHER EDUCATION & RESEARCH')
    title.underline = True
    title.font.color.rgb = RGBColor(0,0,0)
    temp = heading.add_run('''\n(A DEEMED TO BE UNIVERSITY)''')
    temp.italic = True
    temp.font.color.rgb = RGBColor(0x00,0x00,0x00)
    exam = heading.add_run('''\n'''+ str(paperConfig['exam name']))
    exam.font.color.rgb = RGBColor(0,0,0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    rules = document.add_paragraph()
    time = rules.add_run('Time:' + str(paperConfig['Max duration']) + '\t\t\t\t\t\t\t\t\tMax Marks:' + paperConfig['Max marks'])
    #time = rules.add_run('Time: 3 hours')
    time.alignment = WD_ALIGN_PARAGRAPH.LEFT
    time.bold = True
    time.font.size = Pt(7)
    
    subject = document.add_heading(level = 3)
    sub = subject.add_run(paperConfig['subject'])
    sub.font.color.rgb = RGBColor(0,0,0)
    subject.alignment = WD_ALIGN_PARAGRAPH.CENTER
   
    important_notes = document.add_paragraph()
    
    temp = paperConfig['ins'].split('.')
    res = '\n'
    for sent in temp:
        res = res + sent + '\n'

    i = important_notes.add_run(res)

    i.italic = True
    i.font.size = Pt(8)
    important_notes.alignment = WD_ALIGN_PARAGRAPH.CENTER


    cat1  = document.add_heading(level = 4)
    topic = cat1.add_run('\nLONG ESSAY')
    topic.font.color.rgb = RGBColor(0,0,0)
    topic.font.size = Pt(10.5)
    topic.underline = True
    topic = cat1.add_run('\t\t\t\t\t\t2 X 10 = 20 Marks')
    topic.font.color.rgb = RGBColor(0,0,0)
    topic.font.size = Pt(10.5)
    result = df[df['category'] == 'le']
    for a,b in result.iterrows():
        p = document.add_paragraph(style = 'ListNumber') 
        q = p.add_run(b['question'])
        q.font.size = Pt(8)
    
    cat2  = document.add_heading(level = 4)
    topic = cat2.add_run('\nSHORT ESSAY')
    topic.font.color.rgb = RGBColor(0,0,0)
    topic.font.size = Pt(10.5)
    topic.underline = True
    topic = cat2.add_run('\t\t\t\t\t\t10 X 5 = 50 Marks')
    topic.font.color.rgb = RGBColor(0,0,0)
    topic.font.size = Pt(10.5)
    result = df[df['category'] == 'se']
    for a,b in result.iterrows():
        p = document.add_paragraph(style = 'ListNumber') 
        q = p.add_run(b['question'])
        q.font.size = Pt(8)

    cat3  = document.add_heading(level = 4)
    topic = cat3.add_run('\nSHORT ANSWERS')
    topic.font.color.rgb = RGBColor(0,0,0)
    topic.font.size = Pt(10.5)
    topic.underline = True
    topic = cat3.add_run('\t\t\t\t\t\t10 X 3 = 30 Marks')
    topic.font.color.rgb = RGBColor(0,0,0)
    topic.font.size = Pt(10.5)
    result = df[df['category'] == 'sa']
    for a,b in result.iterrows():
        p = document.add_paragraph(style = 'ListNumber') 
        q = p.add_run(b['question'])
        q.font.size = Pt(8)
        
    document.save('first_paper.docx')
    return True


def setPaper_with_options(df, paperConfig, name):
    document = Document()

    q_code = document.add_heading('Question Paper Code:' + str(paperConfig['Question Paper Code']),level = 6)
    q_code.alignment = WD_ALIGN_PARAGRAPH.RIGHT


    heading = document.add_heading(level = 3)
    title = heading.add_run('SRI DEVRAJ URS ACADEMY OF HIGHER EDUCATION & RESEARCH')
    title.underline = True
    title.font.color.rgb = RGBColor(0,0,0)
    temp = heading.add_run('''\n(A DEEMED TO BE UNIVERSITY)''')
    temp.italic = True
    temp.font.color.rgb = RGBColor(0x00,0x00,0x00)
    exam = heading.add_run('''\n'''+ str(paperConfig['exam name']))
    exam.font.color.rgb = RGBColor(0,0,0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    rules = document.add_paragraph()
    time = rules.add_run('Time:' + str(paperConfig['Max duration']) + '\t\t\t\t\t\t\t\t\tMax Marks:' + paperConfig['Max marks'])
    #time = rules.add_run('Time: 3 hours')
    time.alignment = WD_ALIGN_PARAGRAPH.LEFT
    time.bold = True
    time.font.size = Pt(7)
    
    subject = document.add_heading(level = 3)
    sub = subject.add_run(paperConfig['subject'])
    sub.font.color.rgb = RGBColor(0,0,0)
    subject.alignment = WD_ALIGN_PARAGRAPH.CENTER
   
    important_notes = document.add_paragraph()
    temp = paperConfig['ins'].split('.')
    res = '\n'
    for sent in temp:
        res = res + sent + '\n'

    i = important_notes.add_run(res)

    

    i.italic = True
    i.font.size = Pt(8)
    important_notes.alignment = WD_ALIGN_PARAGRAPH.CENTER


    cat1  = document.add_heading(level = 4)
    topic = cat1.add_run('\nLONG ESSAY ( Answer any 2 )')
    topic.font.color.rgb = RGBColor(0,0,0)
    topic.font.size = Pt(10.5)
    topic.underline = True
    topic = cat1.add_run('\t\t\t\t2 X 10 = 20 Marks')
    topic.font.color.rgb = RGBColor(0,0,0)
    topic.font.size = Pt(10.5)
    result = df[df['category'] == 'le']
    for a,b in result.iterrows():
        p = document.add_paragraph(style = 'ListNumber') 
        q = p.add_run(b['question'])
        q.font.size = Pt(8)
    '''
    document.add_paragraph(
        'first item in ordered list', style='ListNumber'
    )

    document.add_paragraph(
        'second item in ordered list', style='ListNumber'
    )
    '''

    cat2  = document.add_heading(level = 4)
    topic = cat2.add_run('\nSHORT ESSAY ( Answer any 10 )')
    topic.font.color.rgb = RGBColor(0,0,0)
    topic.font.size = Pt(10.5)
    topic.underline = True
    topic = cat2.add_run('\t\t\t10 X 5 = 50 Marks')
    topic.font.color.rgb = RGBColor(0,0,0)
    topic.font.size = Pt(10.5)
    result = df[df['category'] == 'se']
    for a,b in result.iterrows():
        p = document.add_paragraph(style = 'ListNumber') 
        q = p.add_run(b['question'])
        q.font.size = Pt(8)

    cat3  = document.add_heading(level = 4)
    topic = cat3.add_run('\nSHORT ANSWERS ( No choices )')
    topic.font.color.rgb = RGBColor(0,0,0)
    topic.font.size = Pt(10.5)
    topic.underline = True
    topic = cat3.add_run('\t\t\t10 X 3 = 30 Marks')
    topic.font.color.rgb = RGBColor(0,0,0)
    topic.font.size = Pt(10.5)
    result = df[df['category'] == 'sa']
    for a,b in result.iterrows():
        p = document.add_paragraph(style = 'ListNumber') 
        q = p.add_run(b['question'])
        q.font.size = Pt(8)
        
    document.save('optionsPaper.docx')
    return True
